"""
make_docx.py
============
Convert manuscript/manuscript.md into an editable Word document
(manuscript/manuscript.docx) for submission/editing.

Handles: headings (#/##/###), paragraphs with **bold** and *italic* inline
markup, pipe tables (header row bold), fenced code blocks (monospace), bullet
lists, and horizontal rules.

Usage: .venv/bin/python analysis/make_docx.py
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "manuscript" / "manuscript.md"
DST = ROOT / "manuscript" / "manuscript.docx"


def add_runs(paragraph, text):
    """Add runs to a paragraph, honouring **bold** and *italic* inline markup."""
    # tokenize: **bold** first, then *italic* within non-bold segments
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            subparts = re.split(r"(\*.+?\*)", part)
            for sp in subparts:
                if not sp:
                    continue
                if sp.startswith("*") and sp.endswith("*"):
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sp)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_table(rows):
    """rows: list of raw markdown table lines; returns (header, body)."""
    data = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        data.append(cells)
    # drop separator row(s)
    data = [d for d in data if not all(re.fullmatch(r":?-{3,}:?", c) for c in d)]
    return data[0], data[1:]


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    title_done = False
    while i < n:
        line = lines[i]

        # horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # fenced code block
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # closing fence
            p = doc.add_paragraph()
            p.style = doc.styles["No Spacing"]
            run = p.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and not title_done:
                p = doc.add_heading(text, level=0)
                title_done = True
            else:
                doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # table block
        if is_table_row(line):
            block = []
            while i < n and is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            header, body = parse_table(block)
            ncols = len(header)
            table = doc.add_table(rows=1, cols=ncols)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for c, htext in enumerate(header):
                hdr[c].text = ""
                run = hdr[c].paragraphs[0].add_run(htext)
                run.bold = True
            for row in body:
                cells = table.add_row().cells
                for c in range(ncols):
                    val = row[c] if c < len(row) else ""
                    val = val.replace("**", "")
                    cells[c].text = val
            doc.add_paragraph()
            continue

        # bullet list
        if re.match(r"^\s*-\s+", line):
            text = re.sub(r"^\s*-\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text)
            i += 1
            continue

        # blank
        if not line.strip():
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_runs(p, line.strip())
        i += 1

    doc.save(DST)
    print("Saved:", DST)


if __name__ == "__main__":
    main()
