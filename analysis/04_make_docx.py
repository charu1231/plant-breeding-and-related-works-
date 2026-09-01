#!/usr/bin/env python3
"""04_make_docx.py — render paper/PAPER.md into a real Word document (PAPER.docx)
with proper headings, tables, bold/italic runs and embedded figures."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
MD = (ROOT / "paper" / "PAPER.md").read_text(encoding="utf-8")
OUT = ROOT / "paper" / "PAPER.docx"

doc = Document()
# ---- base styles
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)
for h, sz in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)]:
    s = doc.styles[h]; s.font.size = Pt(sz); s.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)

def add_runs(par, text):
    """Parse **bold**, *italic*, and `code` inline markdown into docx runs."""
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"
        else:
            par.add_run(tok)

def add_table(rows):
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    hdr, body = rows[0], rows[1:]   # markdown separator row was already dropped
    t = doc.add_table(rows=1 + len(body), cols=ncol)
    t.style = "Table Grid"
    for j, cell in enumerate(hdr):
        p = t.rows[0].cells[j].paragraphs[0]
        add_runs(p, cell)
        for r in p.runs: r.bold = True; r.font.size = Pt(9)
        t.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(body, start=1):
        for j, cell in enumerate(row):
            p = t.rows[i].cells[j].paragraphs[0]
            add_runs(p, cell)
            for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

lines = MD.splitlines()
i = 0
# ---- YAML-ish front matter -> title block
if lines[0].strip() == "---":
    i = 1
    fm = {}
    key = None
    while i < len(lines) and lines[i].strip() != "---":
        ln = lines[i]
        m = re.match(r"^(\w[\w-]*):\s*(.*)$", ln)
        if m:
            key, val = m.group(1), m.group(2).strip().strip('"')
            fm[key] = val
        elif key and ln.startswith(" "):
            fm[key] += " " + ln.strip().strip('"')
        i += 1
    i += 1
    if "title" in fm:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(fm["title"].replace("|", "").strip()); r.bold = True; r.font.size = Pt(16)
    for k, sz, it in [("authors", 11, False), ("date", 10, False), ("status", 10, True)]:
        if k in fm:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(fm[k]); r.font.size = Pt(sz); r.italic = it
    doc.add_paragraph()

while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    if not s:
        i += 1; continue
    if s == "---":
        i += 1; continue
    m = re.match(r"^(#{1,3})\s+(.*)$", s)
    if m:
        doc.add_heading(re.sub(r"\*\*", "", m.group(2)), level=len(m.group(1)))
        i += 1; continue
    im = re.match(r"^!\[(.*?)\]\((.*?)\)", s)
    if im:
        alt, path = im.groups()
        img = (ROOT / "paper" / path).resolve()
        if not img.exists():
            img = (ROOT / path).resolve()
        if img.exists():
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img), width=Inches(6.3))
        i += 1; continue
    if s.startswith("|"):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(set(c) <= set("-: ") for c in cells):
                tbl.append(cells)
            i += 1
        if len(tbl) >= 2:
            add_table(tbl)
        continue
    if re.match(r"^[-*]\s+", s):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r"^[-*]\s+", "", s))
        i += 1; continue
    if re.match(r"^\d+\.\s+", s):
        p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s+", "", s))
        i += 1; continue
    # paragraph: merge consecutive non-special lines
    buf = [s]
    i += 1
    while i < len(lines):
        nxt = lines[i].strip()
        if (not nxt or nxt.startswith(("#", "|", "!", "-", "*", ">")) or nxt == "---"
                or re.match(r"^\d+\.\s+", nxt)):
            break
        buf.append(nxt); i += 1
    p = doc.add_paragraph(); add_runs(p, " ".join(buf))

doc.save(OUT)
print("wrote", OUT, f"({OUT.stat().st_size/1e6:.2f} MB)")
